"""Unit tests for DOCXRedactor run-level text substitution and formatting preservation."""

import tempfile
import unittest
from pathlib import Path

import docx

from src.models import PIIEntity, SourceLocation, SourceType
from src.redaction.docx_redactor import DOCXRedactor


class TestDOCXRedactor(unittest.TestCase):
    """Test suite for DOCXRedactor handling single run, multi-run, and table cell redactions."""

    def setUp(self) -> None:
        """Create a synthetic DOCX file with formatted runs and tables."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.docx_path = Path(self.temp_dir.name) / "synthetic_redact.docx"

        doc = docx.Document()

        # Paragraph 1: Single run PII with bold styling
        p1 = doc.add_paragraph()
        r1_1 = p1.add_run("Contact ")
        r1_2 = p1.add_run("john@example.com")
        r1_2.bold = True
        r1_3 = p1.add_run(" for help.")

        # Paragraph 2: Multi-run split PII (john@ inside run1, example.com inside run2)
        p2 = doc.add_paragraph()
        r2_1 = p2.add_run("Send mail to john@")
        r2_2 = p2.add_run("example.com")
        r2_2.italic = True
        r2_3 = p2.add_run(" today.")

        # Table 1: PII inside table cell
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "User Name"
        table.cell(0, 1).text = "Email"
        table.cell(1, 0).text = "Rashi Patil"
        table.cell(1, 1).text = "rashi.patil@gmail.com"

        doc.save(str(self.docx_path))

    def tearDown(self) -> None:
        """Clean up temporary directory."""
        self.temp_dir.cleanup()

    def test_single_run_redaction(self) -> None:
        """Test redacting PII contained entirely within one run."""
        loc = SourceLocation(source_type=SourceType.PARAGRAPH, paragraph_index=0)
        entity = PIIEntity(
            entity_type="EMAIL",
            original_text="john@example.com",
            start=8,
            end=24,
            confidence=1.0,
            detector="email_detector",
            replacement="user@example.com",
            source_location=loc,
        )

        redactor = DOCXRedactor(self.docx_path)
        out_path = Path(self.temp_dir.name) / "out_single.docx"
        summary = redactor.redact_document([entity], out_path)

        self.assertGreaterEqual(summary.get("_TOTAL_REPLACED", 0), 1)

        # Reopen output document and verify text
        res_doc = docx.Document(str(out_path))
        res_p = res_doc.paragraphs[0]
        self.assertIn("user@example.com", res_p.text)
        self.assertNotIn("john@example.com", res_p.text)

    def test_multi_run_split_redaction(self) -> None:
        """Test redacting PII split across multiple runs."""
        loc = SourceLocation(source_type=SourceType.PARAGRAPH, paragraph_index=1)
        # "Send mail to john@example.com today."
        # "john@example.com" spans 13..29
        entity = PIIEntity(
            entity_type="EMAIL",
            original_text="john@example.com",
            start=13,
            end=29,
            confidence=1.0,
            detector="email_detector",
            replacement="john.doe@example.com",
            source_location=loc,
        )

        redactor = DOCXRedactor(self.docx_path)
        out_path = Path(self.temp_dir.name) / "out_multi.docx"
        redactor.redact_document([entity], out_path)

        res_doc = docx.Document(str(out_path))
        res_p = res_doc.paragraphs[1]
        self.assertIn("john.doe@example.com", res_p.text)
        self.assertNotIn("john@example.com", res_p.text)

    def test_table_cell_redaction(self) -> None:
        """Test redacting PII inside a table cell."""
        loc = SourceLocation(
            source_type=SourceType.TABLE_CELL,
            table_index=0,
            row_index=1,
            cell_index=1,
            paragraph_index=0,
        )
        entity = PIIEntity(
            entity_type="EMAIL",
            original_text="rashi.patil@gmail.com",
            start=0,
            end=21,
            confidence=1.0,
            detector="email_detector",
            replacement="rashi@example.com",
            source_location=loc,
        )

        redactor = DOCXRedactor(self.docx_path)
        out_path = Path(self.temp_dir.name) / "out_table.docx"
        redactor.redact_document([entity], out_path)

        res_doc = docx.Document(str(out_path))
        res_cell = res_doc.tables[0].cell(1, 1)
        self.assertIn("rashi@example.com", res_cell.text)
        self.assertNotIn("rashi.patil@gmail.com", res_cell.text)


if __name__ == "__main__":
    unittest.main()
