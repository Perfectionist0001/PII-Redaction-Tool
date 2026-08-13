"""Unit tests for DOCXExtractor using synthetic DOCX files."""

import tempfile
import unittest
from pathlib import Path

import docx

from src.extractors.docx_extractor import DOCXExtractor
from src.models import SourceType


class TestDOCXExtractor(unittest.TestCase):
    """Tests for extracting paragraphs, tables, and runs from DOCX files."""

    def setUp(self) -> None:
        """Create a synthetic DOCX document for testing."""
        self.temp_dir = tempfile.TemporaryDirectory()
        self.docx_path = Path(self.temp_dir.name) / "synthetic_test.docx"

        doc = docx.Document()

        # Add paragraph
        p = doc.add_paragraph()
        run1 = p.add_run("John Smith ")
        run2 = p.add_run("john@example.com")

        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Email"
        table.cell(1, 0).text = "John Smith"
        table.cell(1, 1).text = "john@example.com"

        doc.save(str(self.docx_path))

    def tearDown(self) -> None:
        """Clean up temporary test directory."""
        self.temp_dir.cleanup()

    def test_synthetic_docx_extraction(self) -> None:
        """Verify extraction of paragraphs and table content from synthetic DOCX."""
        extractor = DOCXExtractor(self.docx_path)
        chunks = extractor.extract_chunks(include_empty=False)

        # Check total extracted non-empty chunks (1 paragraph + 4 table cells)
        self.assertGreaterEqual(len(chunks), 5)

        # Filter paragraph chunks
        para_chunks = [c for c in chunks if c.source_location.source_type == SourceType.PARAGRAPH]
        self.assertEqual(len(para_chunks), 1)
        self.assertEqual(para_chunks[0].text, "John Smith john@example.com")
        self.assertEqual(len(para_chunks[0].runs), 2)
        self.assertEqual(para_chunks[0].runs[0].text, "John Smith ")
        self.assertEqual(para_chunks[0].runs[0].start_pos, 0)
        self.assertEqual(para_chunks[0].runs[0].end_pos, 11)
        self.assertEqual(para_chunks[0].runs[1].text, "john@example.com")
        self.assertEqual(para_chunks[0].runs[1].start_pos, 11)

        # Filter table chunks
        table_chunks = [c for c in chunks if c.source_location.source_type == SourceType.TABLE_CELL]
        self.assertEqual(len(table_chunks), 4)

        extracted_texts = [c.text for c in table_chunks]
        self.assertIn("Name", extracted_texts)
        self.assertIn("Email", extracted_texts)
        self.assertIn("John Smith", extracted_texts)
        self.assertIn("john@example.com", extracted_texts)

        # Verify source location mapping
        sample_cell = [c for c in table_chunks if c.text == "john@example.com"][0]
        self.assertEqual(sample_cell.source_location.table_index, 0)
        self.assertEqual(sample_cell.source_location.row_index, 1)
        self.assertEqual(sample_cell.source_location.cell_index, 1)

    def test_summary_generation(self) -> None:
        """Verify diagnostic summary dictionary generation."""
        extractor = DOCXExtractor(self.docx_path)
        summary = extractor.get_summary()

        self.assertEqual(summary["top_paragraphs_count"], 1)
        self.assertEqual(summary["tables_count"], 1)
        self.assertEqual(summary["table_cells_count"], 4)
        self.assertGreater(summary["total_extracted_characters"], 0)


if __name__ == "__main__":
    unittest.main()
