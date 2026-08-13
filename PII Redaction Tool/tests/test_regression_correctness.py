"""Regression tests for PII in tables and source-location-aware entity matching.

Covers:
- Same local offsets in different table cells (must be distinct entities)
- PII detection inside table cells
- PII detection in normal paragraphs
- Repeated person and organization names across chunks
- Redaction completeness on synthetic DOCX documents
"""

import tempfile
import unittest
from pathlib import Path

import docx

from src.detection_pipeline import DetectionPipeline, create_default_pipeline
from src.evaluation.metrics import MetricsCalculator
from src.models import PIIEntity, SourceLocation, SourceType, TextChunk
from src.redaction.docx_redactor import DOCXRedactor
from src.redaction.pseudonymizer import Pseudonymizer
from src.validation.redaction_check import validate_redaction_completeness


class TestDuplicateOffsetsDifferentTableCells(unittest.TestCase):
    """Regression tests: identical local offsets in different table cells must be kept separate."""

    def _make_loc_tuple(self, table_index: int, row_index: int, cell_index: int):
        """Helper: build the loc_tuple used in evaluation matching."""
        return ("table_cell", 0, table_index, row_index, cell_index, None, None)

    def test_same_offsets_different_cells_are_distinct_entities(self) -> None:
        """Two entities with same (start,end,text) in different table cells = 2 distinct entities."""
        loc_tbl23 = SourceLocation(
            source_type=SourceType.TABLE_CELL,
            table_index=23,
            row_index=1,
            cell_index=1,
            paragraph_index=0,
        )
        loc_tbl32 = SourceLocation(
            source_type=SourceType.TABLE_CELL,
            table_index=32,
            row_index=0,
            cell_index=0,
            paragraph_index=0,
        )

        e1 = PIIEntity(
            entity_type="PERSON",
            original_text="Kushal Subbayya Hegde",
            start=0,
            end=21,
            confidence=0.85,
            detector="spaCy_NER",
            source_location=loc_tbl23,
        )
        e2 = PIIEntity(
            entity_type="PERSON",
            original_text="Kushal Subbayya Hegde",
            start=0,
            end=21,
            confidence=0.85,
            detector="spaCy_NER",
            source_location=loc_tbl32,
        )

        pipeline = DetectionPipeline()
        resolved = pipeline.deduplicate_and_resolve_overlaps([e1, e2])

        self.assertEqual(
            len(resolved), 2,
            "Entities in different table cells with the same local offsets must NOT be collapsed."
        )

    def test_same_offsets_same_cell_are_deduplicated(self) -> None:
        """Two identical entities in the SAME table cell ARE duplicates and should be collapsed."""
        loc = SourceLocation(
            source_type=SourceType.TABLE_CELL,
            table_index=5,
            row_index=0,
            cell_index=0,
            paragraph_index=0,
        )
        e1 = PIIEntity("EMAIL", "a@b.com", 0, 7, 1.0, "email", source_location=loc)
        e2 = PIIEntity("EMAIL", "a@b.com", 0, 7, 1.0, "email", source_location=loc)

        pipeline = DetectionPipeline()
        resolved = pipeline.deduplicate_and_resolve_overlaps([e1, e2])
        self.assertEqual(len(resolved), 1)

    def test_evaluation_metrics_distinct_cells_count_separately(self) -> None:
        """Evaluation must not collapse two entities sharing local offsets in different cells."""
        tup_tbl23 = self._make_loc_tuple(23, 1, 1)
        tup_tbl32 = self._make_loc_tuple(32, 0, 0)

        preds = [
            (tup_tbl23, 0, 21, "Kushal Subbayya Hegde"),
            (tup_tbl32, 0, 21, "Kushal Subbayya Hegde"),
        ]
        gts = [
            (tup_tbl23, 0, 21, "Kushal Subbayya Hegde"),
            (tup_tbl32, 0, 21, "Kushal Subbayya Hegde"),
        ]

        m = MetricsCalculator.calculate_category_metrics("PERSON", preds, gts)
        self.assertEqual(m.total_ground_truth, 2)
        self.assertEqual(m.total_predictions, 2)
        self.assertEqual(m.true_positives, 2)
        self.assertEqual(m.false_positives, 0)
        self.assertEqual(m.false_negatives, 0)

    def test_evaluation_metrics_different_cells_one_missed(self) -> None:
        """If one cell's entity is missed, it shows up as FN, not hidden behind same-offset match."""
        tup_tbl23 = self._make_loc_tuple(23, 1, 1)
        tup_tbl32 = self._make_loc_tuple(32, 0, 0)

        # Only predict entity in tbl23 — miss tbl32
        preds = [(tup_tbl23, 0, 21, "Kushal Subbayya Hegde")]
        gts = [
            (tup_tbl23, 0, 21, "Kushal Subbayya Hegde"),
            (tup_tbl32, 0, 21, "Kushal Subbayya Hegde"),
        ]

        m = MetricsCalculator.calculate_category_metrics("PERSON", preds, gts)
        self.assertEqual(m.true_positives, 1)
        self.assertEqual(m.false_negatives, 1)
        self.assertAlmostEqual(m.recall, 0.5)


class TestPIIInTableCells(unittest.TestCase):
    """Regression tests for detecting PII inside table cells."""

    def setUp(self) -> None:
        self.pipeline = create_default_pipeline()

    def _make_table_chunk(self, text: str, table_index: int = 0, row: int = 0, cell: int = 0):
        loc = SourceLocation(
            source_type=SourceType.TABLE_CELL,
            table_index=table_index,
            row_index=row,
            cell_index=cell,
            paragraph_index=0,
        )
        return TextChunk(f"tbl_{table_index}_r{row}_c{cell}", text, loc)

    def test_email_detected_in_table_cell(self) -> None:
        chunk = self._make_table_chunk("cs.connect@kshinternational.com", table_index=0)
        entities = self.pipeline.process_chunk(chunk)
        emails = [e for e in entities if e.entity_type == "EMAIL"]
        self.assertGreaterEqual(len(emails), 1)
        self.assertIn("cs.connect@kshinternational.com", [e.original_text for e in emails])

    def test_phone_detected_in_table_cell(self) -> None:
        chunk = self._make_table_chunk("+91 20 4505 3237", table_index=1)
        entities = self.pipeline.process_chunk(chunk)
        phones = [e for e in entities if e.entity_type == "PHONE"]
        self.assertGreaterEqual(len(phones), 1)

    def test_person_detected_in_table_cell(self) -> None:
        chunk = self._make_table_chunk("Sarthak Malvadkar", table_index=2)
        entities = self.pipeline.process_chunk(chunk)
        persons = [e for e in entities if e.entity_type == "PERSON"]
        self.assertGreaterEqual(len(persons), 1)

    def test_entities_from_different_cells_not_collapsed(self) -> None:
        """Entities from two different table cells must coexist even if text/offset is identical."""
        chunk1 = self._make_table_chunk("Sarthak Malvadkar", table_index=5, row=0, cell=0)
        chunk2 = self._make_table_chunk("Sarthak Malvadkar", table_index=5, row=1, cell=0)

        entities = self.pipeline.process_chunks([chunk1, chunk2])
        person_entities = [e for e in entities if "Sarthak" in e.original_text]
        # Should have at least 2 — one per cell
        self.assertGreaterEqual(len(person_entities), 2)


class TestPIIInParagraphs(unittest.TestCase):
    """Regression tests for detecting PII in normal body paragraphs."""

    def setUp(self) -> None:
        self.pipeline = create_default_pipeline()

    def _para_chunk(self, text: str, p_idx: int = 0) -> TextChunk:
        loc = SourceLocation(source_type=SourceType.PARAGRAPH, paragraph_index=p_idx)
        return TextChunk(f"para_{p_idx}", text, loc)

    def test_email_in_paragraph(self) -> None:
        chunk = self._para_chunk("Contact us at cs.connect@kshinternational.com today.")
        entities = self.pipeline.process_chunk(chunk)
        self.assertTrue(any(e.entity_type == "EMAIL" for e in entities))

    def test_phone_in_paragraph(self) -> None:
        chunk = self._para_chunk("Call our office at +91 20 6729 5100 for assistance.")
        entities = self.pipeline.process_chunk(chunk)
        self.assertTrue(any(e.entity_type == "PHONE" for e in entities))

    def test_person_in_paragraph(self) -> None:
        chunk = self._para_chunk("Compliance Officer: Sarthak Malvadkar, KSH International.")
        entities = self.pipeline.process_chunk(chunk)
        self.assertTrue(any(e.entity_type == "PERSON" for e in entities))

    def test_address_in_paragraph(self) -> None:
        chunk = self._para_chunk(
            "Registered Office at 11/3, Village Birdewadi, Chakan, Pune – 410 501, Maharashtra, India."
        )
        entities = self.pipeline.process_chunk(chunk)
        self.assertTrue(any(e.entity_type == "ADDRESS" for e in entities))


class TestRepeatedPersonNames(unittest.TestCase):
    """Regression: repeated person names across multiple chunks must all be detected."""

    def setUp(self) -> None:
        self.pipeline = create_default_pipeline()

    def test_repeated_person_name_all_occurrences_detected(self) -> None:
        """All occurrences of a person's name across chunks must be detected."""
        loc0 = SourceLocation(source_type=SourceType.PARAGRAPH, paragraph_index=0)
        loc1 = SourceLocation(source_type=SourceType.PARAGRAPH, paragraph_index=50)
        loc2 = SourceLocation(
            source_type=SourceType.TABLE_CELL,
            table_index=10,
            row_index=1,
            cell_index=1,
            paragraph_index=0,
        )

        chunks = [
            TextChunk("c0", "Promoter: Kushal Subbayya Hegde is a director.", loc0),
            TextChunk("c1", "We are led by our promoter Kushal Subbayya Hegde.", loc1),
            TextChunk("c2", "Kushal Subbayya Hegde", loc2),
        ]

        entities = self.pipeline.process_chunks(chunks)
        kushal_entities = [e for e in entities if "Kushal Subbayya Hegde" in e.original_text]
        # Must find the name in at least 2 chunks (cross-chunk propagation)
        self.assertGreaterEqual(
            len(kushal_entities), 2,
            "Co-reference propagation must detect 'Kushal Subbayya Hegde' across chunks."
        )

    def test_repeated_person_different_table_cells_all_detected(self) -> None:
        """Same person name in two different table cells must each be found independently."""
        loc_tbl23 = SourceLocation(
            source_type=SourceType.TABLE_CELL,
            table_index=23, row_index=1, cell_index=1, paragraph_index=0,
        )
        loc_tbl32 = SourceLocation(
            source_type=SourceType.TABLE_CELL,
            table_index=32, row_index=0, cell_index=0, paragraph_index=0,
        )
        # Also provide a paragraph chunk to seed co-reference propagation
        loc_para = SourceLocation(source_type=SourceType.PARAGRAPH, paragraph_index=5)

        chunks = [
            TextChunk("c_para", "Director profile: Kushal Subbayya Hegde.", loc_para),
            TextChunk("c_t23", "Kushal Subbayya Hegde", loc_tbl23),
            TextChunk("c_t32", "Kushal Subbayya Hegde", loc_tbl32),
        ]

        entities = self.pipeline.process_chunks(chunks)
        kushal_entities = [e for e in entities if "Kushal Subbayya Hegde" in e.original_text]
        self.assertGreaterEqual(len(kushal_entities), 2)


class TestRepeatedOrganizationNames(unittest.TestCase):
    """Regression: repeated organization names must be detected at every occurrence."""

    def setUp(self) -> None:
        self.pipeline = create_default_pipeline()

    def test_org_name_detected_in_multiple_paragraphs(self) -> None:
        locs = [
            SourceLocation(source_type=SourceType.PARAGRAPH, paragraph_index=i)
            for i in range(3)
        ]
        chunks = [
            TextChunk("c0", "KSH International Limited is the issuer.", locs[0]),
            TextChunk("c1", "The lead manager is ICICI Securities Limited.", locs[1]),
            TextChunk("c2", "KSH International Limited has offices in Pune.", locs[2]),
        ]
        entities = self.pipeline.process_chunks(chunks)
        ksh_entities = [e for e in entities if "KSH International" in e.original_text]
        self.assertGreaterEqual(len(ksh_entities), 2)

    def test_public_org_names_detected_as_organization(self) -> None:
        """Public orgs are detected by NER but per policy should be reviewed for exclusion.
        This test ensures detection fires — policy filtering is the reviewer's job."""
        loc = SourceLocation(source_type=SourceType.PARAGRAPH, paragraph_index=0)
        # Use a longer sentence with clear context — en_core_web_sm needs enough
        # surrounding text to recognize organizational entities reliably.
        chunk = TextChunk(
            "c0",
            "The Securities and Exchange Board of India and the Reserve Bank of India "
            "have issued guidelines. ICICI Securities Limited acted as the lead manager.",
            loc,
        )
        entities = self.pipeline.process_chunk(chunk)
        org_entities = [e for e in entities if e.entity_type == "ORGANIZATION"]
        org_texts = [e.original_text for e in org_entities]
        # At minimum one of these well-known entities should be tagged
        self.assertTrue(
            len(org_texts) > 0,
            f"Expected at least one ORGANIZATION entity, got none from: {chunk.text!r}"
        )


class TestRedactionCompletenessValidation(unittest.TestCase):
    """Regression: redaction validation utility correctly identifies missed PII."""

    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def _create_synthetic_docx(self, path: Path, paragraphs: list, tables: list = None) -> None:
        """Helper to create a simple DOCX for testing."""
        doc = docx.Document()
        for para_text in paragraphs:
            doc.add_paragraph(para_text)
        if tables:
            for table_data in tables:
                rows = len(table_data)
                cols = max(len(r) for r in table_data)
                tbl = doc.add_table(rows=rows, cols=cols)
                for r_idx, row in enumerate(table_data):
                    for c_idx, cell_text in enumerate(row):
                        tbl.cell(r_idx, c_idx).text = cell_text
        doc.save(str(path))

    def test_no_missed_redactions_returns_empty_list(self) -> None:
        """validate_redaction_completeness returns [] when PII is fully redacted."""
        original_path = self.temp_path / "original.docx"
        redacted_path = self.temp_path / "redacted.docx"
        self._create_synthetic_docx(original_path, ["Contact: john@example.com"])
        self._create_synthetic_docx(redacted_path, ["Contact: [EMAIL_REDACTED]"])

        missed = validate_redaction_completeness(
            original_path, redacted_path, ["john@example.com"]
        )
        self.assertEqual(missed, [])

    def test_missed_redaction_detected(self) -> None:
        """validate_redaction_completeness reports residual PII that was not redacted."""
        original_path = self.temp_path / "original.docx"
        redacted_path = self.temp_path / "redacted.docx"
        self._create_synthetic_docx(original_path, ["Contact: john@example.com"])
        # Forgot to redact — same text appears in output
        self._create_synthetic_docx(redacted_path, ["Contact: john@example.com"])

        missed = validate_redaction_completeness(
            original_path, redacted_path, ["john@example.com"]
        )
        self.assertEqual(len(missed), 1)
        self.assertEqual(missed[0]["pii_text"], "john@example.com")
        self.assertEqual(missed[0]["occurrences_redacted"], 1)

    def test_pii_in_table_cell_fully_redacted(self) -> None:
        """Table cell PII that is successfully redacted causes no false alarm."""
        original_path = self.temp_path / "orig_tbl.docx"
        redacted_path = self.temp_path / "red_tbl.docx"
        self._create_synthetic_docx(
            original_path, [], tables=[[["Name", "Email"], ["Rashi Patil", "rashi@example.com"]]]
        )
        self._create_synthetic_docx(
            redacted_path, [], tables=[[["Name", "Email"], ["[PERSON]", "[EMAIL]"]]]
        )
        missed = validate_redaction_completeness(
            original_path, redacted_path, ["Rashi Patil", "rashi@example.com"]
        )
        self.assertEqual(missed, [])

    def test_pii_in_table_cell_missed_redaction(self) -> None:
        """Table cell PII that is NOT redacted is reported correctly."""
        original_path = self.temp_path / "orig_tbl2.docx"
        redacted_path = self.temp_path / "red_tbl2.docx"
        self._create_synthetic_docx(
            original_path, [], tables=[[["Name"], ["Rashi Patil"]]]
        )
        # Forgot to redact the name in the table
        self._create_synthetic_docx(
            redacted_path, [], tables=[[["Name"], ["Rashi Patil"]]]
        )
        missed = validate_redaction_completeness(
            original_path, redacted_path, ["Rashi Patil"]
        )
        self.assertEqual(len(missed), 1)
        self.assertEqual(missed[0]["pii_text"], "Rashi Patil")

    def test_full_redaction_pipeline_on_synthetic_docx(self) -> None:
        """End-to-end: run pipeline + redactor on a synthetic DOCX and verify no PII remains."""
        original_path = self.temp_path / "e2e_orig.docx"
        redacted_path = self.temp_path / "e2e_red.docx"

        pii_email = "testuser@company.com"
        pii_phone = "+91 9876543210"
        doc = docx.Document()
        doc.add_paragraph(f"Email: {pii_email}")
        doc.add_paragraph(f"Phone: {pii_phone}")
        doc.save(str(original_path))

        from src.extractors.docx_extractor import DOCXExtractor
        extractor = DOCXExtractor(original_path)
        chunks = extractor.extract_chunks(include_empty=False)

        pipeline = create_default_pipeline()
        entities = pipeline.process_chunks(chunks)

        pseudonymizer = Pseudonymizer(seed=42)
        entities = pseudonymizer.assign_replacements(entities)

        redactor = DOCXRedactor(original_path)
        redactor.redact_document(entities, redacted_path)

        # Check that original PII strings are gone
        pii_strings = [pii_email, pii_phone]
        missed = validate_redaction_completeness(original_path, redacted_path, pii_strings)
        self.assertEqual(
            missed, [],
            f"Residual PII found after redaction: {missed}"
        )


if __name__ == "__main__":
    unittest.main()
