"""Main CLI entrypoint for PII Redaction Tool - Complete Pipeline Orchestrator."""

import argparse
import sys
import time
from pathlib import Path
from typing import Dict, List

import docx

from src.config import settings
from src.detection_pipeline import create_default_pipeline
from src.evaluation.evaluate import run_evaluation
from src.extractors.docx_extractor import DOCXExtractor
from src.models import PIIEntity, TextChunk
from src.redaction.docx_redactor import DOCXRedactor
from src.redaction.pseudonymizer import Pseudonymizer


def build_parser() -> argparse.ArgumentParser:
    """Build CLI argument parser for PII Redaction Tool."""
    parser = argparse.ArgumentParser(
        description="PII Redaction Tool — Microsoft Word (.docx) PII Detection & Pseudonymization Engine",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "--input",
        "-i",
        type=str,
        default=str(settings.primary_input_file),
        help="Path to input Microsoft Word (.docx) document",
    )
    parser.add_argument(
        "--output",
        "-o",
        type=str,
        default=str(settings.output_dir / "redacted_prospectus.docx"),
        help="Target output path for redacted DOCX document",
    )
    parser.add_argument(
        "--evaluate",
        action="store_true",
        help="Run evaluation metrics against ground truth after redaction",
    )
    parser.add_argument(
        "--ground-truth",
        type=str,
        default=str(settings.evaluation_dir / "ground_truth.json"),
        help="Path to ground_truth.json for evaluation",
    )
    parser.add_argument(
        "--report",
        type=str,
        default=str(settings.evaluation_dir / "evaluation_report.md"),
        help="Path to write evaluation_report.md",
    )
    parser.add_argument(
        "--verbose",
        "-v",
        action="store_true",
        help="Enable detailed diagnostic logging",
    )
    return parser


def run_pipeline(
    input_path: Path,
    output_path: Path,
    evaluate: bool = False,
    ground_truth_path: Path | None = None,
    report_path: Path | None = None,
    verbose: bool = False,
) -> int:
    """Execute complete PII detection, pseudonymization, redaction, and optional evaluation pipeline.

    Args:
        input_path: Target input DOCX file path.
        output_path: Target output DOCX file path.
        evaluate: Whether to run evaluation against ground truth.
        ground_truth_path: Ground truth JSON file path.
        report_path: Target path for evaluation_report.md.
        verbose: Verbose logging flag.

    Returns:
        Exit code (0 for success, 1 for error).
    """
    start_time = time.time()
    print("==================================================================")
    print("      PII REDACTION ENGINE — INDUSTRIAL DOCX PROCESSING TOOL       ")
    print("==================================================================")

    # Step 1: Input Validation
    print(f"[1/8] Validating input file: {input_path.name}")
    if not input_path.exists():
        print(f"ERROR: Input file does not exist at {input_path.resolve()}")
        return 1

    if input_path.suffix.lower() != ".docx":
        print(f"ERROR: Unsupported file format '{input_path.suffix}'. Only .docx supported.")
        return 1

    # Step 2: Extract Paragraphs & Tables
    print("[2/8] Extracting paragraphs and table cells from document...")
    try:
        extractor = DOCXExtractor(input_path)
        chunks: List[TextChunk] = extractor.extract_chunks(include_empty=False)
        print(f"      Extracted {len(chunks):,} non-empty text chunks.")
    except Exception as err:
        print(f"ERROR: Extraction failed: {err}")
        return 1

    # Step 3 & 4: Run Detectors & Aggregate
    print("[3/8] Running local PII detection pipeline (Regex + spaCy NER)...")
    try:
        pipeline = create_default_pipeline()
        raw_entities: List[PIIEntity] = pipeline.process_chunks(chunks)
        print(f"      Detected {len(raw_entities):,} finalized non-overlapping PII entities.")
    except Exception as err:
        print(f"ERROR: Detection pipeline failed: {err}")
        return 1

    # Step 5 & 6: Pseudonymization
    print("[5/8] Generating deterministic synthetic pseudonym replacements...")
    try:
        pseudonymizer = Pseudonymizer(seed=42)
        processed_entities: List[PIIEntity] = pseudonymizer.assign_replacements(raw_entities)
        print(f"      Mapped {len(processed_entities):,} entities to deterministic fake values.")
    except Exception as err:
        print(f"ERROR: Pseudonymization failed: {err}")
        return 1

    # Step 7 & 8: Redact DOCX & Save Output
    print(f"[7/8] Applying run-level redactions and writing output: {output_path.name}")
    try:
        redactor = DOCXRedactor(input_path)
        category_counts: Dict[str, int] = redactor.redact_document(
            processed_entities, output_path
        )
        total_replaced = category_counts.pop("_TOTAL_REPLACED", 0)
        print(f"      Successfully replaced {total_replaced:,} entity spans.")
    except Exception as err:
        print(f"ERROR: Redaction failed: {err}")
        return 1

    # Validate output document integrity
    try:
        val_doc = docx.Document(str(output_path))
        if verbose:
            print(
                f"      Output Validation OK: {len(val_doc.paragraphs):,} paragraphs, "
                f"{len(val_doc.tables):,} tables preserved."
            )
    except Exception as err:
        print(f"ERROR: Output file corruption detected: {err}")
        return 1

    # Step 9 & 10: Optional Evaluation
    if evaluate:
        print("\n[9/10] Running quantitative evaluation against ground truth...")
        gt_file = ground_truth_path or (settings.evaluation_dir / "ground_truth.json")
        rep_file = report_path or (settings.evaluation_dir / "evaluation_report.md")

        if not gt_file.exists():
            print(f"ERROR: Ground truth file not found at {gt_file.resolve()}")
            return 1

        try:
            eval_results = run_evaluation(input_path, gt_file, rep_file)
            tok = eval_results.token_metrics
            print("==================================================================")
            print("EVALUATION RESULTS SUMMARY")
            print("==================================================================")
            if getattr(eval_results, "is_partial_review", False):
                print("  [PARTIAL HUMAN-REVIEW MODE - PRECISION/RECALL/F1 NOT AVAILABLE]")
                print(f"  - Reviewed Candidates   : {eval_results.reviewed_candidates}")
                print(f"  - Accepted Ground Truth : {eval_results.total_ground_truth}")
                print(f"  - Rejected Candidates   : {eval_results.rejected_candidates}")
                print(f"  - Known True Positives  : {eval_results.true_positives}")
                print(f"  - Known False Positives : {eval_results.false_positives}")
                print(f"  - Known False Negatives : {eval_results.false_negatives}")
                print(f"  - Unassessed Predictions: {eval_results.unassessed_predictions}")
            else:
                print(f"  - Micro-Precision : {eval_results.overall_precision:.4f} ({eval_results.overall_precision * 100:.2f}%)")
                print(f"  - Micro-Recall    : {eval_results.overall_recall:.4f} ({eval_results.overall_recall * 100:.2f}%)")
                print(f"  - Micro-F1 Score  : {eval_results.overall_f1:.4f} ({eval_results.overall_f1 * 100:.2f}%)")
                print(f"  - Token Accuracy  : {tok.accuracy:.4f} ({tok.accuracy * 100:.2f}%) across {tok.total_tokens:,} tokens")
            print(f"  - Report Saved To : {rep_file.resolve()}")
        except Exception as err:
            print(f"ERROR: Evaluation failed: {err}")
            return 1

    elapsed = time.time() - start_time
    print("\n==================================================================")
    print("EXECUTION SUMMARY REPORT")
    print("==================================================================")
    print(f"Input Document       : {input_path.resolve()}")
    print(f"Redacted Output      : {output_path.resolve()}")
    print(f"Total Detections     : {len(raw_entities):,}")
    print(f"Total Replaced Spans : {total_replaced:,}")
    print("\nReplaced Entities by Category:")
    for cat, cnt in sorted(category_counts.items()):
        print(f"  - {cat:<20}: {cnt:,}")
    print("------------------------------------------------------------------")
    print(f"Status               : SUCCESS")
    print(f"Total Time Elapsed   : {elapsed:.2f}s")
    print("==================================================================")

    return 0


def main() -> None:
    """CLI main entrypoint."""
    parser = build_parser()
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    gt_path = Path(args.ground_truth) if args.ground_truth else None
    
    # Prefer verified ground truth if it exists and default was used
    if gt_path and gt_path.name == "ground_truth.json":
        verified_path = gt_path.parent / "ground_truth_verified.json"
        if verified_path.exists():
            gt_path = verified_path

    rep_path = Path(args.report) if args.report else None

    exit_code = run_pipeline(
        input_path=input_path,
        output_path=output_path,
        evaluate=args.evaluate,
        ground_truth_path=gt_path,
        report_path=rep_path,
        verbose=args.verbose,
    )
    sys.exit(exit_code)


if __name__ == "__main__":
    main()
