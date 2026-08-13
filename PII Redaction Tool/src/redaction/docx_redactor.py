"""DOCX document redaction engine using run-level text substitution."""

from collections import defaultdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import docx
from docx.text.paragraph import Paragraph

from src.models import PIIEntity, SourceType


class DOCXRedactor:
    """Applies finalized PII entity redactions to a DOCX document preserving formatting."""

    def __init__(self, docx_path: Path | str) -> None:
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"Input DOCX file not found: {self.docx_path}")
        self.doc = docx.Document(str(self.docx_path))

    def redact_paragraph(self, paragraph: Paragraph, entities: List[PIIEntity]) -> int:
        """Redact entities within a single paragraph using run-level substitution.

        Args:
            paragraph: python-docx Paragraph object.
            entities: List of PIIEntity objects located in this paragraph.

        Returns:
            Number of entities successfully replaced.
        """
        if not entities or not paragraph.runs:
            return 0

        # Process entities in reverse order of start offset (right-to-left)
        sorted_entities = sorted(entities, key=lambda e: e.start, reverse=True)
        replaced_count = 0

        for entity in sorted_entities:
            replacement = entity.replacement
            if replacement is None:
                continue

            ent_start = entity.start
            ent_end = entity.end

            # Calculate character start/end boundaries for each run
            run_bounds: List[Tuple[int, int]] = []
            curr = 0
            for r in paragraph.runs:
                r_len = len(r.text)
                run_bounds.append((curr, curr + r_len))
                curr += r_len

            # Find overlapping run indices
            first_run_idx = None
            last_run_idx = None

            for r_idx, (r_start, r_end) in enumerate(run_bounds):
                if r_end > ent_start and first_run_idx is None:
                    first_run_idx = r_idx
                if r_start < ent_end:
                    last_run_idx = r_idx

            if first_run_idx is None or last_run_idx is None:
                continue

            # Case A: Entity fits inside a single run
            if first_run_idx == last_run_idx:
                r = paragraph.runs[first_run_idx]
                r_start, _ = run_bounds[first_run_idx]
                l_start = ent_start - r_start
                l_end = ent_end - r_start
                r.text = r.text[:l_start] + replacement + r.text[l_end:]
                replaced_count += 1

            # Case B: Entity is split across multiple runs
            else:
                r_first = paragraph.runs[first_run_idx]
                r1_start, _ = run_bounds[first_run_idx]
                l_start = ent_start - r1_start
                r_first.text = r_first.text[:l_start] + replacement

                # Clear text in intermediate runs
                for k in range(first_run_idx + 1, last_run_idx):
                    paragraph.runs[k].text = ""

                r_last = paragraph.runs[last_run_idx]
                rL_start, _ = run_bounds[last_run_idx]
                l_end = ent_end - rL_start
                r_last.text = r_last.text[l_end:]
                replaced_count += 1

        return replaced_count

    def _get_paragraph_object(self, loc: SourceType | Any) -> Optional[Paragraph]:
        """Retrieve the python-docx Paragraph object for a given SourceLocation."""
        if not loc:
            return None
        s_type = loc.source_type
        p_idx = loc.paragraph_index

        if s_type == SourceType.PARAGRAPH and p_idx is not None:
            if p_idx < len(self.doc.paragraphs):
                return self.doc.paragraphs[p_idx]

        elif (
            s_type == SourceType.TABLE_CELL
            and loc.table_index is not None
            and loc.row_index is not None
            and loc.cell_index is not None
            and p_idx is not None
        ):
            if loc.table_index < len(self.doc.tables):
                tbl = self.doc.tables[loc.table_index]
                if loc.row_index < len(tbl.rows):
                    row = tbl.rows[loc.row_index]
                    if loc.cell_index < len(row.cells):
                        cell = row.cells[loc.cell_index]
                        if p_idx < len(cell.paragraphs):
                            return cell.paragraphs[p_idx]

        elif (
            s_type == SourceType.HEADER
            and loc.header_index is not None
            and p_idx is not None
        ):
            if loc.header_index < len(self.doc.sections):
                header = self.doc.sections[loc.header_index].header
                if header and p_idx < len(header.paragraphs):
                    return header.paragraphs[p_idx]

        elif (
            s_type == SourceType.FOOTER
            and loc.footer_index is not None
            and p_idx is not None
        ):
            if loc.footer_index < len(self.doc.sections):
                footer = self.doc.sections[loc.footer_index].footer
                if footer and p_idx < len(footer.paragraphs):
                    return footer.paragraphs[p_idx]

        return None

    def redact_document(
        self, entities: List[PIIEntity], output_path: Path | str
    ) -> Dict[str, int]:
        """Apply all detections to document paragraphs, tables, and headers, and save output.

        Groups entities by the underlying paragraph XML element to prevent duplicate redaction
        on merged table cells or linked headers/footers.

        Args:
            entities: Finalized list of non-overlapping PIIEntity objects with replacements.
            output_path: Target path to write the redacted DOCX document.

        Returns:
            Dictionary summary of replacement counts by category.
        """
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        category_counts: Dict[str, int] = defaultdict(int)

        # Group entities by the unique paragraph element
        paragraph_entities: Dict[Any, Tuple[Paragraph, List[PIIEntity]]] = {}

        for e in entities:
            if not e.source_location or e.replacement is None:
                continue

            p_obj = self._get_paragraph_object(e.source_location)
            if p_obj is None:
                continue

            category_counts[e.entity_type] += 1
            p_key = p_obj._p  # Unique paragraph XML element reference
            
            if p_key not in paragraph_entities:
                paragraph_entities[p_key] = (p_obj, [])
            paragraph_entities[p_key][1].append(e)

        total_replaced = 0

        # Apply redactions to unique paragraph objects
        for p_obj, p_ents in paragraph_entities.values():
            # Deduplicate entity spans for the same physical paragraph that might
            # have been registered under different coordinates (e.g. merged cells)
            seen_spans = set()
            unique_ents = []
            for ent in p_ents:
                span_key = (ent.start, ent.end, ent.entity_type)
                if span_key not in seen_spans:
                    seen_spans.add(span_key)
                    unique_ents.append(ent)

            total_replaced += self.redact_paragraph(p_obj, unique_ents)

        # Save redacted document to output path
        self.doc.save(str(output_file))

        category_counts["_TOTAL_REPLACED"] = total_replaced
        return dict(category_counts)
