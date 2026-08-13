"""Redaction validation utility checking if detected original PII values remain in the output document.

Validates:
- Paragraphs
- Tables
- Headers
- Footers

Reports:
- Original entity text (summarized/masked for privacy)
- Entity type
- Source location
- Whether it remains in output
- Replacement if known
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

import docx
from docx.text.paragraph import Paragraph

from src.models import PIIEntity, SourceLocation, SourceType


def _mask_pii_safe(text: str) -> str:
    """Safely mask PII text for console logging to prevent privacy leaks."""
    text = text.strip()
    if len(text) <= 3:
        return "*" * len(text)
    return text[0] + "*" * (len(text) - 2) + text[-1]


def _extract_location_text(doc: docx.Document, loc: SourceLocation | Dict[str, Any]) -> Optional[str]:
    """Retrieve the text at a specific source location in a DOCX document."""
    if isinstance(loc, SourceLocation):
        s_type = loc.source_type.value if isinstance(loc.source_type, SourceType) else str(loc.source_type)
        p_idx = loc.paragraph_index
        tbl_idx = loc.table_index
        r_idx = loc.row_index
        c_idx = loc.cell_index
        h_idx = loc.header_index
        f_idx = loc.footer_index
    else:
        s_type = loc.get("source_type", "paragraph")
        p_idx = loc.get("paragraph_index")
        tbl_idx = loc.get("table_index")
        r_idx = loc.get("row_index")
        c_idx = loc.get("cell_index")
        h_idx = loc.get("header_index")
        f_idx = loc.get("footer_index")

    if s_type == "paragraph" and p_idx is not None and p_idx < len(doc.paragraphs):
        return doc.paragraphs[p_idx].text

    if (
        s_type == "table_cell"
        and tbl_idx is not None
        and r_idx is not None
        and c_idx is not None
        and p_idx is not None
    ):
        if tbl_idx < len(doc.tables):
            tbl = doc.tables[tbl_idx]
            if r_idx < len(tbl.rows):
                row = tbl.rows[r_idx]
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    if p_idx < len(cell.paragraphs):
                        return cell.paragraphs[p_idx].text

    if s_type == "header" and h_idx is not None and p_idx is not None:
        if h_idx < len(doc.sections):
            hdr = doc.sections[h_idx].header
            if hdr and p_idx < len(hdr.paragraphs):
                return hdr.paragraphs[p_idx].text

    if s_type == "footer" and f_idx is not None and p_idx is not None:
        if f_idx < len(doc.sections):
            ftr = doc.sections[f_idx].footer
            if ftr and p_idx < len(ftr.paragraphs):
                return ftr.paragraphs[p_idx].text

    return None


def validate_redacted_document(
    redacted_docx_path: Path | str,
    entities: List[PIIEntity] | List[Dict[str, Any]],
    case_sensitive: bool = True,
) -> List[Dict[str, Any]]:
    """Check if detected PII entities remain in the output DOCX at their respective locations or document-wide.

    Args:
        redacted_docx_path: Path to output redacted DOCX.
        entities: List of PIIEntity objects or ground truth entity dictionaries.
        case_sensitive: Whether text checking is case sensitive.

    Returns:
        List of report records for entities that REMAIN in the output.
    """
    path = Path(redacted_docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Redacted DOCX file not found: {path}")

    doc = docx.Document(str(path))
    remaining_report: List[Dict[str, Any]] = []

    for entity in entities:
        if isinstance(entity, PIIEntity):
            orig_text = entity.original_text.strip()
            etype = entity.entity_type
            loc = entity.source_location
            replacement = entity.replacement
        else:
            orig_text = entity.get("text", "").strip()
            etype = entity.get("entity_type", "UNKNOWN")
            loc = entity.get("source_location")
            replacement = entity.get("replacement")

        if not orig_text or len(orig_text) < 2:
            continue

        loc_text = _extract_location_text(doc, loc) if loc else None

        search_text = orig_text if case_sensitive else orig_text.lower()

        remains = False
        if loc_text is not None:
            check_loc_text = loc_text if case_sensitive else loc_text.lower()
            if search_text in check_loc_text:
                remains = True
        else:
            # Fallback to checking full body text if location text is unavailable
            body_text = "\n".join([p.text for p in doc.paragraphs])
            check_body = body_text if case_sensitive else body_text.lower()
            if search_text in check_body:
                remains = True

        if remains:
            remaining_report.append(
                {
                    "original_entity": orig_text,
                    "safe_masked_text": _mask_pii_safe(orig_text),
                    "entity_type": etype,
                    "source_location": str(loc) if loc else "Unknown",
                    "remains_in_output": True,
                    "replacement": replacement,
                }
            )

    return remaining_report


def main() -> int:
    """CLI entrypoint for redaction validation."""
    parser = argparse.ArgumentParser(description="DOCX Redaction Completeness Validator")
    parser.add_argument("--redacted", required=True, help="Path to redacted output DOCX")
    parser.add_argument("--ground-truth", required=True, help="Path to ground_truth.json")
    parser.add_argument("--case-insensitive", action="store_true", help="Case-insensitive check")

    args = parser.parse_args()
    redacted_path = Path(args.redacted)
    gt_path = Path(args.ground_truth)

    if not gt_path.exists():
        print(f"Error: Ground truth file not found: {gt_path}", file=sys.stderr)
        return 1

    with open(gt_path, "r", encoding="utf-8") as f:
        gt_data = json.load(f)

    gt_entities = gt_data.get("annotated_entities", [])
    print(f"Loaded {len(gt_entities):,} ground-truth entities for validation.")
    print(f"Validating document: {redacted_path.name}...")

    failures = validate_redacted_document(
        redacted_path,
        gt_entities,
        case_sensitive=not args.case_insensitive,
    )

    if not failures:
        print("VALIDATION SUCCESSFUL: No target PII entities remain in the output document.")
        return 0

    print(f"VALIDATION FAILURE: {len(failures)} entity occurrence(s) remain unredacted in output.")
    # Safe summarized logging (masked PII to prevent console exposure)
    summary_by_type: Dict[str, int] = {}
    for item in failures:
        etype = item["entity_type"]
        summary_by_type[etype] = summary_by_type.get(etype, 0) + 1

    print("Unredacted counts by category:", summary_by_type)
    print("\nFirst 10 sample unredacted locations (masked text):")
    for item in failures[:10]:
        print(f"  - [{item['entity_type']}] Masked: {item['safe_masked_text']} at Location: {item['source_location']}")

    return 1


if __name__ == "__main__":
    sys.exit(main())
