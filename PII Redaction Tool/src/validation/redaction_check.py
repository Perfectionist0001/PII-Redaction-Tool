"""Redaction validation utility — checks that a redacted DOCX contains no residual PII.

Usage (CLI):
    python -m src.validation.redaction_check \\
        --original "input/Red Herring Prospectus.docx" \\
        --redacted "output/redacted_prospectus.docx" \\
        --ground-truth "evaluation/ground_truth.json"

Usage (Python / pytest):
    from src.validation.redaction_check import validate_redaction_completeness
    missed = validate_redaction_completeness(original_path, redacted_path, pii_strings)
    assert missed == [], f"Missed redactions: {missed}"
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

import docx


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------


def _extract_all_text(docx_path: Path) -> str:
    """Extract all visible text from a DOCX including paragraphs, tables, headers, footers."""
    doc = docx.Document(str(docx_path))
    parts: List[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        parts.append(para.text)

    # Table cells (all paragraphs in every cell)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    parts.append(para.text)

    # Section headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                parts.append(para.text)
        if section.footer:
            for para in section.footer.paragraphs:
                parts.append(para.text)

    return "\n".join(parts)


def _extract_pii_strings_from_ground_truth(gt_path: Path) -> List[str]:
    """Load all PII text values from a ground_truth.json file."""
    with open(gt_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    entities: List[Dict[str, Any]] = data.get("annotated_entities", [])
    # Deduplicate preserving order
    seen = set()
    result: List[str] = []
    for e in entities:
        t = e.get("text", "").strip()
        if t and t not in seen:
            seen.add(t)
            result.append(t)
    return result


# ---------------------------------------------------------------------------
# Core validation function
# ---------------------------------------------------------------------------


def validate_redaction_completeness(
    original_path: Path,
    redacted_path: Path,
    pii_strings: List[str],
    *,
    case_sensitive: bool = True,
) -> List[Dict[str, Any]]:
    """Check the redacted document for residual occurrences of known PII strings.

    Args:
        original_path:  Path to the original (pre-redaction) DOCX.
        redacted_path:  Path to the redacted output DOCX.
        pii_strings:    List of PII text values to search for.
        case_sensitive: Whether the search is case-sensitive (default True).

    Returns:
        List of dicts, one per missed redaction:
        {
            "pii_text": str,          # The PII string found in the redacted doc
            "occurrences_original": int,  # Count in original
            "occurrences_redacted": int,  # Count remaining in redacted (should be 0)
        }
        Returns an empty list if no residual PII is found.
    """
    if not original_path.exists():
        raise FileNotFoundError(f"Original DOCX not found: {original_path}")
    if not redacted_path.exists():
        raise FileNotFoundError(f"Redacted DOCX not found: {redacted_path}")

    original_text = _extract_all_text(original_path)
    redacted_text = _extract_all_text(redacted_path)

    if not case_sensitive:
        original_text_cmp = original_text.lower()
        redacted_text_cmp = redacted_text.lower()
    else:
        original_text_cmp = original_text
        redacted_text_cmp = redacted_text

    missed: List[Dict[str, Any]] = []

    for pii in pii_strings:
        if not pii or len(pii.strip()) < 3:
            continue  # Skip trivially short strings that generate false matches

        search_pii = pii if case_sensitive else pii.lower()

        count_original = original_text_cmp.count(search_pii)
        count_redacted = redacted_text_cmp.count(search_pii)

        if count_original == 0:
            # The PII didn't actually appear in the original — skip
            continue

        if count_redacted > 0:
            missed.append(
                {
                    "pii_text": pii,
                    "occurrences_original": count_original,
                    "occurrences_redacted": count_redacted,
                }
            )

    return missed


# ---------------------------------------------------------------------------
# CLI entrypoint
# ---------------------------------------------------------------------------


def main() -> int:
    """CLI entrypoint for redaction completeness validation."""
    parser = argparse.ArgumentParser(
        description="Validate that a redacted DOCX contains no residual PII from ground truth."
    )
    parser.add_argument(
        "--original",
        required=True,
        help="Path to the original (pre-redaction) DOCX",
    )
    parser.add_argument(
        "--redacted",
        required=True,
        help="Path to the redacted output DOCX",
    )
    parser.add_argument(
        "--ground-truth",
        required=True,
        help="Path to ground_truth.json containing known PII strings",
    )
    parser.add_argument(
        "--case-insensitive",
        action="store_true",
        help="Perform case-insensitive search (default: case-sensitive)",
    )
    args = parser.parse_args()

    original = Path(args.original)
    redacted = Path(args.redacted)
    gt_path = Path(args.ground_truth)

    if not gt_path.exists():
        print(f"ERROR: Ground truth file not found: {gt_path}", file=sys.stderr)
        return 1

    print(f"Loading PII strings from: {gt_path.name}")
    pii_strings = _extract_pii_strings_from_ground_truth(gt_path)
    print(f"  {len(pii_strings):,} unique PII strings to check.")

    print(f"Validating redacted document: {redacted.name}")
    missed = validate_redaction_completeness(
        original,
        redacted,
        pii_strings,
        case_sensitive=not args.case_insensitive,
    )

    if not missed:
        print("PASS - No residual PII strings found in the redacted document.")
        return 0

    print(f"\nFAIL - {len(missed)} PII string(s) still present in the redacted document:\n")
    for item in missed:
        print(
            f"  [{item['pii_text']!r}]  "
            f"original={item['occurrences_original']}x  "
            f"redacted={item['occurrences_redacted']}x  <- NOT FULLY REDACTED"
        )
    return 1


if __name__ == "__main__":
    sys.exit(main())
