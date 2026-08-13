"""DOCX document text and structure extractor using python-docx."""

from pathlib import Path
from typing import Any, Dict, List, Optional

import docx
from docx.text.paragraph import Paragraph

from src.models import RunInfo, SourceLocation, SourceType, TextChunk


class DOCXExtractor:
    """Extracts text chunks, run boundaries, and source metadata from DOCX files."""

    def __init__(self, file_path: Path | str) -> None:
        self.file_path = Path(file_path)
        self._doc: Optional[docx.Document] = None

    def _load_document(self) -> docx.Document:
        """Load and cache the python-docx Document object."""
        if self._doc is None:
            if not self.file_path.exists():
                raise FileNotFoundError(f"DOCX file not found: {self.file_path}")
            self._doc = docx.Document(str(self.file_path))
        return self._doc

    def _extract_paragraph_chunk(
        self,
        paragraph: Paragraph,
        chunk_id: str,
        source_location: SourceLocation,
    ) -> TextChunk:
        """Extract a single paragraph into a TextChunk with run position information."""
        full_text = paragraph.text
        runs: List[RunInfo] = []
        current_pos = 0

        for r_idx, run in enumerate(paragraph.runs):
            run_text = run.text
            start_pos = current_pos
            end_pos = current_pos + len(run_text)
            runs.append(
                RunInfo(
                    run_index=r_idx,
                    text=run_text,
                    start_pos=start_pos,
                    end_pos=end_pos,
                    docx_run=run,
                )
            )
            current_pos = end_pos

        return TextChunk(
            chunk_id=chunk_id,
            text=full_text,
            source_location=source_location,
            runs=runs,
            docx_element=paragraph,
            metadata={"run_count": len(runs), "char_count": len(full_text)},
        )

    def extract_paragraphs(self, include_empty: bool = False) -> List[TextChunk]:
        """Extract all top-level body paragraphs."""
        doc = self._load_document()
        chunks: List[TextChunk] = []

        for p_idx, paragraph in enumerate(doc.paragraphs):
            if not include_empty and not paragraph.text.strip():
                continue
            loc = SourceLocation(
                source_type=SourceType.PARAGRAPH,
                paragraph_index=p_idx,
            )
            chunk_id = f"paragraph_{p_idx}"
            chunks.append(self._extract_paragraph_chunk(paragraph, chunk_id, loc))

        return chunks

    def extract_tables(self, include_empty: bool = False) -> List[TextChunk]:
        """Extract all table cells across all document tables."""
        doc = self._load_document()
        chunks: List[TextChunk] = []

        for tbl_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        if not include_empty and not paragraph.text.strip():
                            continue
                        loc = SourceLocation(
                            source_type=SourceType.TABLE_CELL,
                            table_index=tbl_idx,
                            row_index=r_idx,
                            cell_index=c_idx,
                            paragraph_index=p_idx,
                        )
                        chunk_id = f"tbl_{tbl_idx}_r_{r_idx}_c_{c_idx}_p_{p_idx}"
                        chunk = self._extract_paragraph_chunk(paragraph, chunk_id, loc)
                        chunk.docx_element = cell  # Reference to cell container
                        chunks.append(chunk)

        return chunks

    def extract_headers_and_footers(self, include_empty: bool = False) -> List[TextChunk]:
        """Extract text chunks from section headers and footers."""
        doc = self._load_document()
        chunks: List[TextChunk] = []

        for s_idx, section in enumerate(doc.sections):
            # Headers
            if section.header:
                for p_idx, paragraph in enumerate(section.header.paragraphs):
                    if not include_empty and not paragraph.text.strip():
                        continue
                    loc = SourceLocation(
                        source_type=SourceType.HEADER,
                        header_index=s_idx,
                        paragraph_index=p_idx,
                    )
                    chunk_id = f"header_{s_idx}_p_{p_idx}"
                    chunks.append(self._extract_paragraph_chunk(paragraph, chunk_id, loc))

            # Footers
            if section.footer:
                for p_idx, paragraph in enumerate(section.footer.paragraphs):
                    if not include_empty and not paragraph.text.strip():
                        continue
                    loc = SourceLocation(
                        source_type=SourceType.FOOTER,
                        footer_index=s_idx,
                        paragraph_index=p_idx,
                    )
                    chunk_id = f"footer_{s_idx}_p_{p_idx}"
                    chunks.append(self._extract_paragraph_chunk(paragraph, chunk_id, loc))

        return chunks

    def extract_chunks(self, include_empty: bool = False) -> List[TextChunk]:
        """Extract all text chunks from paragraphs, tables, headers, and footers."""
        chunks: List[TextChunk] = []
        chunks.extend(self.extract_paragraphs(include_empty=include_empty))
        chunks.extend(self.extract_tables(include_empty=include_empty))
        chunks.extend(self.extract_headers_and_footers(include_empty=include_empty))
        return chunks

    def get_summary(self) -> Dict[str, Any]:
        """Generate a diagnostic summary of the document elements and counts."""
        doc = self._load_document()

        top_paragraphs = len(doc.paragraphs)
        tables_count = len(doc.tables)

        cell_count = 0
        table_paragraph_count = 0

        for table in doc.tables:
            for row in table.rows:
                cell_count += len(row.cells)
                for cell in row.cells:
                    table_paragraph_count += len(cell.paragraphs)

        body_chunks = self.extract_paragraphs(include_empty=True)
        table_chunks = self.extract_tables(include_empty=True)
        header_footer_chunks = self.extract_headers_and_footers(include_empty=True)

        total_extracted_chars = sum(len(c.text) for c in body_chunks + table_chunks + header_footer_chunks)

        return {
            "file_name": self.file_path.name,
            "top_paragraphs_count": top_paragraphs,
            "table_paragraphs_count": table_paragraph_count,
            "total_body_paragraphs_count": top_paragraphs + table_paragraph_count,
            "tables_count": tables_count,
            "table_cells_count": cell_count,
            "header_footer_chunks_count": len(header_footer_chunks),
            "total_extracted_characters": total_extracted_chars,
        }
